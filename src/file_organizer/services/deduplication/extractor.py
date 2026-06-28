# pyre-ignore-all-errors
"""Document text extraction module.

Extracts text content from various document formats for semantic analysis.
Supports PDF, DOCX, TXT, RTF, ODT, and Markdown document formats.
"""

from __future__ import annotations

import logging
import os
import sys
import zipfile
from collections.abc import Iterator
from contextlib import contextmanager
from pathlib import Path
from typing import BinaryIO

from file_organizer.utils.safedir import SafeDir

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Extracts text content from various document formats.

    Supported formats:
    - PDF (.pdf)
    - Word Documents (.docx)
    - Plain Text (.txt, .md)
    - Rich Text (.rtf)
    - OpenDocument (.odt)
    """

    def __init__(self) -> None:
        """Initialize the document extractor."""
        self.supported_extensions = {".pdf", ".docx", ".txt", ".rtf", ".odt", ".md"}
        self._check_dependencies()

    @staticmethod
    def _open_reader_fd(file_path: Path, scan_root: Path | None) -> int | None:
        """Return an ``O_NOFOLLOW`` reader fd for *file_path*, or ``None`` for legacy.

        POSIX only:

        - With *scan_root*: anchored traversal —
          ``SafeDir.open_root(scan_root).open_anchored_reader(relative)`` opens
          every intermediate ancestor with ``O_NOFOLLOW``, so a symlink swapped
          into an *ancestor* directory between enumeration and extraction is
          refused (``SymlinkRejected``), closing the nested-ancestor TOCTOU
          (#286/#1269). ``file_path`` must live under *scan_root*; otherwise
          ``relative_to`` raises ``ValueError`` (the per-format / ``extract_text``
          handlers catch it → ``""``), refusing the read rather than falling back.
        - Without *scan_root*: parent-rooted —
          ``SafeDir.open_root(parent).open_for_reader(name)`` (leaf ``O_NOFOLLOW``).

        Returns ``None`` only when SafeDir is unavailable (``NotImplementedError``
        — non-POSIX), signalling the caller to use a plain ``open``. A symlinked
        root/leaf/ancestor (``SymlinkRejected``) or missing path propagates as
        ``OSError`` — it is *not* downgraded to a legacy open.
        """
        if scan_root is not None:
            # Enforce the trusted-root boundary on EVERY platform, *before* any
            # SafeDir-availability fallback: a file outside scan_root (or one
            # whose relative path contains ``..``) is refused even on Windows,
            # where SafeDir is unavailable and the caller would otherwise fall
            # back to a plain open() and ingest out-of-root content (#1269).
            relative = file_path.relative_to(scan_root)  # ValueError if outside → refuse
            if any(part == ".." for part in relative.parts):
                raise ValueError(f"path escapes scan_root via '..': {file_path!r}")
            if sys.platform == "win32":  # pragma: no cover - platform skip
                # SafeDir anchored open is POSIX-only, so the caller will use a
                # plain open() here — which follows symlinks/junctions. The
                # lexical check above doesn't catch a symlinked *ancestor* under
                # scan_root, so verify the *resolved* path still lives under the
                # resolved root before signalling the legacy fallback; refuse an
                # out-of-root reparse target. (Best-effort: resolve() is racy,
                # but it's the only containment primitive without SafeDir.)
                try:
                    file_path.resolve().relative_to(scan_root.resolve())
                except (OSError, RuntimeError, ValueError) as exc:
                    raise ValueError(
                        f"path escapes scan_root after resolution: {file_path!r}"
                    ) from exc
                return None
            try:
                root_cm = SafeDir.open_root(scan_root)
            except NotImplementedError:  # pragma: no cover - platform skip
                return None
            with root_cm as sd:
                return sd.open_anchored_reader(relative)
        if sys.platform == "win32":  # pragma: no cover - platform skip
            return None
        try:
            root_cm = SafeDir.open_root(file_path.parent)
        except NotImplementedError:  # pragma: no cover - platform skip
            return None
        with root_cm as sd:
            return sd.open_for_reader(file_path.name)

    @staticmethod
    @contextmanager
    def _open_binary(file_path: Path, scan_root: Path | None = None) -> Iterator[BinaryIO]:
        """Yield a binary handle for *file_path*, opened through SafeDir on POSIX.

        When *scan_root* is supplied the read is anchored to that trusted root
        (nested-ancestor symlink-safe); otherwise it is parent-rooted
        (leaf symlink-safe). Falls back to a plain ``open(..., "rb")`` only
        where SafeDir is unavailable (Windows). See :meth:`_open_reader_fd`.
        """
        # ``_open_reader_fd`` is called *before* the yield so a SafeDir refusal
        # (SymlinkRejected/OSError) or a ``..``-escape (ValueError) surfaces to
        # the per-format handler rather than being mistaken for the
        # SafeDir-unavailable case (which would double-yield → RuntimeError).
        fd = DocumentExtractor._open_reader_fd(file_path, scan_root)
        if fd is None:
            with open(file_path, "rb") as handle:  # noqa: safedir-required, atomic-write  # content extractor — path is a user file read for hash/content analysis
                yield handle
            return
        # Close the raw fd if fdopen itself fails (e.g. EMFILE under fd
        # exhaustion); once it returns, ``with handle`` owns the close.
        try:
            handle = os.fdopen(fd, "rb", closefd=True)
        except OSError:
            os.close(fd)
            raise
        with handle:
            yield handle

    def extract_text(self, file_path: Path, *, scan_root: Path | None = None) -> str:
        """Extract text from a single document.

        Args:
            file_path: Path to document file
            scan_root: Optional trusted root the file was discovered under.
                When supplied, the read uses SafeDir anchored traversal so a
                symlinked ancestor between enumeration and extraction is
                refused (nested-ancestor TOCTOU, #1269). ``None`` keeps the
                parent-rooted leaf-safe behaviour.

        Returns:
            Extracted text content

        Raises:
            ValueError: If file format not supported
            OSError: If file cannot be read
        """
        if not file_path.exists():
            raise OSError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()

        if not self.supports_format(file_path):
            raise ValueError(f"Unsupported format: {extension}")

        try:
            if extension == ".pdf":
                return self._extract_pdf(file_path, scan_root=scan_root)
            elif extension == ".docx":
                return self._extract_docx(file_path, scan_root=scan_root)
            elif extension == ".txt" or extension == ".md":
                return self._extract_text(file_path, scan_root=scan_root)
            elif extension == ".rtf":
                return self._extract_rtf(file_path, scan_root=scan_root)
            elif extension == ".odt":
                return self._extract_odt(file_path, scan_root=scan_root)
            else:
                logger.warning(f"No extractor for {extension}, treating as text")
                return self._extract_text(file_path, scan_root=scan_root)

        except (OSError, ValueError, ImportError) as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""

    def extract_batch(
        self, file_paths: list[Path], *, scan_root: Path | None = None
    ) -> dict[Path, str]:
        """Extract text from multiple documents in batch.

        Args:
            file_paths: List of document paths
            scan_root: Optional trusted root forwarded to :meth:`extract_text`
                for anchored (nested-ancestor-safe) reads (#1269).

        Returns:
            Dictionary mapping file paths to extracted text
        """
        results = {}

        for file_path in file_paths:
            try:
                text = self.extract_text(file_path, scan_root=scan_root)
                results[file_path] = text
                logger.debug(f"Extracted {len(text)} chars from {file_path.name}")
            except (OSError, ValueError, ImportError) as e:
                logger.warning(f"Failed to extract {file_path}: {e}")
                results[file_path] = ""

        logger.info(f"Batch extracted text from {len(results)} documents")

        return results

    def supports_format(self, file_path: Path) -> bool:
        """Check if a file format is supported.

        Args:
            file_path: Path to file

        Returns:
            True if format is supported
        """
        return file_path.suffix.lower() in self.supported_extensions

    def get_supported_formats(self) -> list[str]:
        """Get list of supported file formats.

        Returns:
            List of supported extensions
        """
        return sorted(self.supported_extensions)

    def _extract_pdf(self, file_path: Path, *, scan_root: Path | None = None) -> str:
        """Extract text from PDF file.

        Args:
            file_path: Path to PDF file
            scan_root: Optional trusted root for anchored reads (#1269).

        Returns:
            Extracted text
        """
        try:
            import pypdf

            pypdf_error: type[Exception] = Exception
            try:
                from pypdf.errors import PyPdfError
            except (ImportError, AttributeError):
                pass
            else:
                pypdf_error = PyPdfError

            text_parts = []

            with self._open_binary(file_path, scan_root) as f:
                pdf_reader = pypdf.PdfReader(f)

                # Extract text from each page
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

            full_text = "\n".join(text_parts)
            logger.debug(f"Extracted {len(full_text)} chars from PDF: {file_path.name}")

            return full_text

        except ImportError:
            logger.error("pypdf not installed. Install with: pip install pypdf")
            return ""
        except (pypdf_error, OSError, ValueError, KeyError, IndexError) as e:
            logger.error(f"Error extracting PDF {file_path}: {e}")
            return ""

    def _extract_docx(self, file_path: Path, *, scan_root: Path | None = None) -> str:
        """Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file
            scan_root: Optional trusted root for anchored reads (#1269).

        Returns:
            Extracted text
        """
        try:
            import docx

            with self._open_binary(file_path, scan_root) as f:
                doc = docx.Document(f)

            # Extract text from all paragraphs
            text_parts = [paragraph.text for paragraph in doc.paragraphs]

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)

            full_text = "\n".join(text_parts)
            logger.debug(f"Extracted {len(full_text)} chars from DOCX: {file_path.name}")

            return full_text

        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return ""
        except (OSError, ValueError, KeyError, zipfile.BadZipFile) as e:
            # BadZipFile: python-docx raises it when handed a non-OOXML stream
            # (the SafeDir fileobj of a corrupt/non-docx file).
            logger.error(f"Error extracting DOCX {file_path}: {e}")
            return ""

    def _extract_text(self, file_path: Path, *, scan_root: Path | None = None) -> str:
        """Extract text from plain text file.

        Args:
            file_path: Path to text file
            scan_root: Optional trusted root for anchored reads (#1269).

        Returns:
            File contents
        """
        try:
            # Single SafeDir-guarded read, then try multiple encodings on the
            # raw bytes (avoids re-opening the file per encoding).
            with self._open_binary(file_path, scan_root) as f:
                raw = f.read()

            for encoding in ("utf-8", "latin-1", "cp1252", "ascii"):
                try:
                    text = raw.decode(encoding)
                    logger.debug(f"Read {len(text)} chars from text file: {file_path.name}")
                    return text
                except UnicodeDecodeError:
                    continue

            # If all encodings fail, decode with errors='ignore'
            return raw.decode("utf-8", errors="ignore")

        except OSError as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return ""

    def _extract_rtf(self, file_path: Path, *, scan_root: Path | None = None) -> str:
        """Extract text from RTF file.

        Args:
            file_path: Path to RTF file
            scan_root: Optional trusted root for anchored reads (#1269).

        Returns:
            Extracted text
        """
        try:
            # Try using striprtf if available
            try:
                from striprtf.striprtf import rtf_to_text

                with self._open_binary(file_path, scan_root) as f:
                    rtf_content = f.read().decode("utf-8", errors="ignore")

                text = str(rtf_to_text(rtf_content))
                logger.debug(f"Extracted {len(text)} chars from RTF: {file_path.name}")

                return text

            except ImportError:
                # Fallback: simple RTF stripping
                logger.warning("striprtf not installed, using basic extraction")

                with self._open_binary(file_path, scan_root) as f:
                    content = f.read().decode("utf-8", errors="ignore")

                # Very basic RTF stripping (removes control words)
                import re

                text = re.sub(r"\\[a-z]+\d*\s?", "", content)
                text = re.sub(r"[{}]", "", text)
                text = text.strip()

                return text

        except (OSError, ValueError, ImportError) as e:
            logger.error(f"Error extracting RTF {file_path}: {e}")
            return ""

    def _extract_odt(self, file_path: Path, *, scan_root: Path | None = None) -> str:
        """Extract text from ODT file.

        Args:
            file_path: Path to ODT file
            scan_root: Optional trusted root for anchored reads (#1269).

        Returns:
            Extracted text
        """
        try:
            # ``content.xml`` comes from an untrusted document; use defusedxml to
            # reject entity-expansion / external-entity (XXE) payloads that the
            # stdlib parser would otherwise process.
            from defusedxml.ElementTree import ParseError, fromstring as _xml_fromstring

            # ODT files are ZIP archives; open through SafeDir (symlink-safe).
            with self._open_binary(file_path, scan_root) as f, zipfile.ZipFile(f, "r") as odt_zip:
                # Extract content.xml
                content_xml = odt_zip.read("content.xml")

            # Parse XML (defused against XXE / billion-laughs)
            root = _xml_fromstring(content_xml)

            # Extract all text nodes
            # ODT uses OpenDocument namespace
            namespaces = {
                "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
                "office": "urn:oasis:names:tc:opendocument:xmlns:office:1.0",
            }

            text_parts = []

            # Find all text:p (paragraph) elements
            for paragraph in root.findall(".//text:p", namespaces):
                if paragraph.text:
                    text_parts.append(paragraph.text)

                # Also get text from child elements
                for child in paragraph:
                    if child.text:
                        text_parts.append(child.text)
                    if child.tail:
                        text_parts.append(child.tail)

            full_text = "\n".join(text_parts)
            logger.debug(f"Extracted {len(full_text)} chars from ODT: {file_path.name}")

            return full_text

        except (OSError, KeyError, ValueError, zipfile.BadZipFile, ParseError, ImportError) as e:
            # ParseError (malformed content.xml — a SyntaxError subclass) and
            # ImportError (defusedxml unavailable) are handled here so a bad or
            # untrusted ODT degrades to "" rather than crashing the dedup flow.
            logger.error(f"Error extracting ODT {file_path}: {e}")
            return ""

    def _check_dependencies(self) -> None:
        """Check if required dependencies are installed."""
        dependencies = {
            "pypdf": "PDF extraction",
            "docx": "DOCX extraction",
        }

        missing = []
        for module, purpose in dependencies.items():
            try:
                __import__(module if module != "docx" else "docx")
            except ImportError:
                missing.append(f"{module} ({purpose})")

        if missing:
            logger.warning(
                f"Missing optional dependencies: {', '.join(missing)}. "
                "Some document formats may not be supported."
            )
